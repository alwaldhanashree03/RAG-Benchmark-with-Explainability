"""Data loader for MS MARCO and Natural Questions datasets.

Industry best practice: Use official datasets via Hugging Face datasets library.
Reference: MS MARCO paper (Bajaj et al., 2016)
"""

from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import json
import io

import pandas as pd
from datasets import load_dataset
from loguru import logger
from tqdm import tqdm

from src.utils.config_loader import get_config

# Import PDF and DOCX parsers with fallback
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        logger.warning("PDF support not available. Install pypdf or PyPDF2 to enable PDF uploads.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("DOCX support not available. Install python-docx to enable DOCX uploads.")


class DatasetLoader:
    """Load and preprocess QA datasets for RAG benchmarking.
    
    Best practice: Use cached datasets to avoid repeated downloads.
    """

    def __init__(self):
        """Initialize dataset loader."""
        self.config = get_config()
        self.dataset_name = self.config.get("dataset.name", "msmarco")
        self.cache_dir = Path(self.config.get("dataset.cache_dir", "./data/raw"))
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def load_msmarco(
        self,
        num_queries: Optional[int] = None,
        num_passages: Optional[int] = None,
    ) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Load MS MARCO passage ranking dataset.
        
        Args:
            num_queries: Limit number of queries (None for all)
            num_passages: Limit number of passages (None for all)
            
        Returns:
            Tuple of (queries_df, passages_df)
            
        Reference:
            MS MARCO: A Human Generated MAchine Reading COmprehension Dataset
            https://microsoft.github.io/msmarco/
        """
        logger.info("Loading MS MARCO dataset...")
        
        num_queries = num_queries or self.config.get("dataset.num_queries", 500)
        num_passages = num_passages or self.config.get("dataset.num_passages", 10000)
        
        try:
            # Load passages (corpus)
            logger.info(f"Loading MS MARCO passages (limit: {num_passages})...")
            passages_dataset = load_dataset(
                "ms_marco",
                "v2.1",
                split=f"train[:{num_passages}]",
                cache_dir=str(self.cache_dir),
            )
            
            passages_data = []
            passage_id_counter = 0
            
            for item in tqdm(passages_dataset, desc="Processing passages"):
                # MS MARCO v2.1 has passages nested in 'passages' field
                passages_list = item.get("passages", {})
                passage_texts = passages_list.get("passage_text", [])
                
                # Extract each passage text
                for passage_text in passage_texts:
                    if passage_text and passage_text.strip():
                        passages_data.append({
                            "passage_id": f"passage_{passage_id_counter}",
                            "text": passage_text.strip(),
                        })
                        passage_id_counter += 1
                        
                        # Stop if we've reached the limit
                        if len(passages_data) >= num_passages:
                            break
                
                if len(passages_data) >= num_passages:
                    break
            
            passages_df = pd.DataFrame(passages_data)
            logger.info(f"Loaded {len(passages_df)} passages")
            
            # Load queries
            logger.info(f"Loading MS MARCO queries (limit: {num_queries})...")
            queries_dataset = load_dataset(
                "ms_marco",
                "v2.1",
                split=f"validation[:{num_queries}]",
                cache_dir=str(self.cache_dir),
            )
            
            queries_data = []
            for item in tqdm(queries_dataset, desc="Processing queries"):
                # Get relevant passage IDs if available
                relevant_passages = item.get("passages", {}).get("passage_id", [])
                
                queries_data.append({
                    "query_id": item.get("query_id", item.get("id")),
                    "query": item.get("query", ""),
                    "answers": item.get("answers", []),
                    "relevant_passage_ids": relevant_passages,
                })
            
            queries_df = pd.DataFrame(queries_data)
            logger.info(f"Loaded {len(queries_df)} queries")
            
            return queries_df, passages_df
            
        except Exception as e:
            logger.error(f"Error loading MS MARCO dataset: {e}")
            logger.info("Falling back to sample data generation...")
            return self._generate_sample_data(num_queries, num_passages)

    def load_natural_questions(
        self,
        num_samples: Optional[int] = None,
    ) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Load Natural Questions dataset.
        
        Args:
            num_samples: Number of samples to load
            
        Returns:
            Tuple of (queries_df, passages_df)
            
        Reference:
            Natural Questions: A Benchmark for Question Answering Research
            https://ai.google.com/research/NaturalQuestions
        """
        logger.info("Loading Natural Questions dataset...")
        
        num_samples = num_samples or self.config.get("dataset.num_queries", 500)
        
        try:
            # Load NQ dataset
            dataset = load_dataset(
                "natural_questions",
                split=f"validation[:{num_samples}]",
                cache_dir=str(self.cache_dir),
            )
            
            queries_data = []
            passages_data = []
            
            for idx, item in enumerate(tqdm(dataset, desc="Processing NQ")):
                question = item.get("question", {}).get("text", "")
                document_text = item.get("document", {}).get("text", "")
                
                # Extract answer
                annotations = item.get("annotations", [{}])[0]
                short_answers = annotations.get("short_answers", [])
                answer_text = ""
                if short_answers:
                    answer_text = short_answers[0].get("text", "")
                
                queries_data.append({
                    "query_id": f"nq_{idx}",
                    "query": question,
                    "answers": [answer_text] if answer_text else [],
                    "relevant_passage_ids": [f"nq_passage_{idx}"],
                })
                
                passages_data.append({
                    "passage_id": f"nq_passage_{idx}",
                    "text": document_text,
                })
            
            queries_df = pd.DataFrame(queries_data)
            passages_df = pd.DataFrame(passages_data)
            
            logger.info(f"Loaded {len(queries_df)} queries and {len(passages_df)} passages")
            
            return queries_df, passages_df
            
        except Exception as e:
            logger.error(f"Error loading Natural Questions: {e}")
            logger.info("Falling back to sample data generation...")
            return self._generate_sample_data(num_samples, num_samples)

    def _generate_sample_data(
        self,
        num_queries: int,
        num_passages: int,
    ) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Generate sample data for testing when real datasets unavailable.
        
        Args:
            num_queries: Number of sample queries
            num_passages: Number of sample passages
            
        Returns:
            Tuple of (queries_df, passages_df)
        """
        logger.warning("Generating sample data for testing purposes")
        
        # Sample passages about AI/ML topics
        sample_passages = [
            "Machine learning is a subset of artificial intelligence that enables systems to learn from data.",
            "Deep learning uses neural networks with multiple layers to process complex patterns.",
            "Natural language processing allows computers to understand and generate human language.",
            "Retrieval-Augmented Generation combines retrieval and generation for better QA systems.",
            "Transformer models revolutionized NLP with attention mechanisms.",
        ]
        
        passages_data = []
        for i in range(min(num_passages, len(sample_passages) * 100)):
            passages_data.append({
                "passage_id": f"sample_passage_{i}",
                "text": sample_passages[i % len(sample_passages)] + f" (Variant {i})",
            })
        
        passages_df = pd.DataFrame(passages_data)
        
        # Sample queries
        sample_queries = [
            {"query": "What is machine learning?", "answer": "A subset of AI that learns from data"},
            {"query": "How does deep learning work?", "answer": "Uses multi-layer neural networks"},
            {"query": "What is NLP?", "answer": "Natural language processing for text understanding"},
            {"query": "What is RAG?", "answer": "Retrieval-Augmented Generation for QA"},
            {"query": "What are transformers?", "answer": "Neural networks with attention mechanisms"},
        ]
        
        queries_data = []
        for i in range(min(num_queries, len(sample_queries) * 100)):
            q = sample_queries[i % len(sample_queries)]
            queries_data.append({
                "query_id": f"sample_query_{i}",
                "query": q["query"] + f" (Variant {i})",
                "answers": [q["answer"]],
                "relevant_passage_ids": [f"sample_passage_{i % len(sample_passages)}"],
            })
        
        queries_df = pd.DataFrame(queries_data)
        
        logger.info(f"Generated {len(queries_df)} sample queries and {len(passages_df)} passages")
        
        return queries_df, passages_df

    def load(self) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Load dataset based on configuration.
        
        Returns:
            Tuple of (queries_df, passages_df)
        """
        if self.dataset_name == "msmarco":
            return self.load_msmarco()
        elif self.dataset_name == "natural_questions":
            return self.load_natural_questions()
        else:
            raise ValueError(f"Unknown dataset: {self.dataset_name}")

    def load_custom_documents(
        self,
        documents: Union[List[str], List[Dict[str, str]]],
        queries: Optional[List[str]] = None,
    ) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Load custom documents and optional queries.
        
        Args:
            documents: List of document texts or dicts with 'text' and optional 'metadata'
            queries: Optional list of query strings
            
        Returns:
            Tuple of (queries_df, passages_df)
        """
        logger.info(f"Loading {len(documents)} custom documents...")
        
        # Process documents
        passages_data = []
        for i, doc in enumerate(documents):
            if isinstance(doc, str):
                passages_data.append({
                    "passage_id": f"custom_passage_{i}",
                    "text": doc.strip(),
                })
            elif isinstance(doc, dict):
                passages_data.append({
                    "passage_id": doc.get("passage_id", f"custom_passage_{i}"),
                    "text": doc.get("text", "").strip(),
                })
        
        passages_df = pd.DataFrame(passages_data)
        logger.info(f"Loaded {len(passages_df)} custom passages")
        
        # Process queries
        if queries:
            queries_data = []
            for i, query in enumerate(queries):
                if query and query.strip():
                    queries_data.append({
                        "query_id": f"custom_query_{i}",
                        "query": query.strip(),
                    })
            queries_df = pd.DataFrame(queries_data)
            logger.info(f"Loaded {len(queries_df)} custom queries")
        else:
            # Create empty queries dataframe
            queries_df = pd.DataFrame(columns=["query_id", "query"])
            logger.info("No queries provided")
        
        return queries_df, passages_df

    def load_from_file(
        self,
        file_path: Union[str, Path],
        file_type: str = "auto",
    ) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Load documents from a file (TXT, JSON, CSV, JSONL, PDF, or DOCX).
        
        Supported formats:
        - TXT: One document per line (or entire file as one document)
        - JSON: {"documents": [...], "queries": [...]} or {"passages": [...], "queries": [...]}
        - JSONL: One JSON object per line with 'text' field
        - CSV: Columns 'text' (required) and optional 'passage_id', 'query'
        - PDF: Extracts text from all pages
        - DOCX: Extracts text from all paragraphs
        
        Args:
            file_path: Path to the file
            file_type: File type ('txt', 'json', 'jsonl', 'csv', 'pdf', 'docx', or 'auto' to detect)
            
        Returns:
            Tuple of (queries_df, passages_df)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Auto-detect file type
        if file_type == "auto":
            file_type = file_path.suffix.lower().lstrip('.')
        
        logger.info(f"Loading data from {file_path} (type: {file_type})...")
        
        documents = []
        queries = []
        
        if file_type == "txt":
            # Load text file - each line is a document
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                # Split by double newline for paragraphs, or single newline for lines
                if '\n\n' in content:
                    documents = [doc.strip() for doc in content.split('\n\n') if doc.strip()]
                else:
                    documents = [line.strip() for line in content.split('\n') if line.strip()]
        
        elif file_type == "json":
            # Load JSON file
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Support multiple JSON structures
            if isinstance(data, dict):
                documents = data.get('documents', data.get('passages', []))
                queries = data.get('queries', [])
            elif isinstance(data, list):
                # Assume it's a list of documents
                documents = data
        
        elif file_type == "jsonl":
            # Load JSONL file
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        obj = json.loads(line)
                        if isinstance(obj, dict):
                            if 'text' in obj or 'passage_text' in obj:
                                documents.append(obj)
                            elif 'query' in obj:
                                queries.append(obj.get('query'))
                        elif isinstance(obj, str):
                            documents.append(obj)
        
        elif file_type == "csv":
            # Load CSV file
            df = pd.read_csv(file_path)
            
            # Check for text column
            text_col = None
            for col in ['text', 'passage', 'document', 'content']:
                if col in df.columns:
                    text_col = col
                    break
            
            if text_col is None:
                raise ValueError(f"CSV must have a 'text', 'passage', 'document', or 'content' column")
            
            documents = df[text_col].tolist()
            
            # Check for queries
            if 'query' in df.columns:
                queries = df['query'].dropna().tolist()
        
        elif file_type == "pdf":
            # Load PDF file
            if not PDF_AVAILABLE:
                raise ImportError("PDF support not available. Install pypdf or PyPDF2: pip install pypdf")
            
            try:
                reader = PdfReader(file_path)
                logger.info(f"PDF has {len(reader.pages)} pages")
                
                # Extract text from each page
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        # Split by paragraphs (double newlines) or use entire page
                        if '\n\n' in text:
                            page_docs = [doc.strip() for doc in text.split('\n\n') if doc.strip()]
                            documents.extend(page_docs)
                        else:
                            documents.append(text.strip())
                
                logger.info(f"Extracted {len(documents)} text segments from PDF")
            except Exception as e:
                raise ValueError(f"Error reading PDF file: {e}")
        
        elif file_type == "docx":
            # Load DOCX file
            if not DOCX_AVAILABLE:
                raise ImportError("DOCX support not available. Install python-docx: pip install python-docx")
            
            try:
                doc = Document(file_path)
                logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
                
                # Extract text from each paragraph
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        documents.append(text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text:
                                documents.append(text)
                
                logger.info(f"Extracted {len(documents)} text segments from DOCX")
            except Exception as e:
                raise ValueError(f"Error reading DOCX file: {e}")
        
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Use 'txt', 'json', 'jsonl', 'csv', 'pdf', or 'docx'")
        
        logger.info(f"Loaded {len(documents)} documents and {len(queries)} queries from file")
        
        return self.load_custom_documents(documents, queries if queries else None)

    def save_processed(
        self,
        queries_df: pd.DataFrame,
        passages_df: pd.DataFrame,
        output_dir: Optional[Path] = None,
    ) -> None:
        """Save processed data to disk.
        
        Args:
            queries_df: Queries dataframe
            passages_df: Passages dataframe
            output_dir: Output directory (default: data/processed)
        """
        output_dir = output_dir or Path("./data/processed")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        queries_path = output_dir / "queries.parquet"
        passages_path = output_dir / "passages.parquet"
        
        queries_df.to_parquet(queries_path, index=False)
        passages_df.to_parquet(passages_path, index=False)
        
        logger.info(f"Saved processed data to {output_dir}")
